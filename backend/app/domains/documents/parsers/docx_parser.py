"""DOCX Word Document Parser with Style-Aware Hierarchy & Table Extraction."""

import io
from typing import List, Optional
import docx
from app.core.security import sanitize_document_text
from app.domains.documents.parsers.base import ParsedChunk, ParsedDocument


class DocxParser:
    """Extracts structured text, headings hierarchy, and tables from DOCX documents."""

    async def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse raw DOCX bytes into a structured ParsedDocument."""
        stream = io.BytesIO(file_bytes)
        doc = docx.Document(stream)

        chunks: List[ParsedChunk] = []
        full_text_parts: List[str] = []
        global_char_offset = 0
        chunk_idx = 0
        current_section = "General Information"
        table_count = 0

        # Estimate page count (roughly 500 words / 3000 chars per page in DOCX)
        total_chars = 0

        for element in doc.element.body:
            # Paragraph
            if element.tag.endswith("p"):
                # Find matching docx paragraph
                p_elem = None
                for p in doc.paragraphs:
                    if p._element == element:
                        p_elem = p
                        break
                if not p_elem:
                    continue

                text = p_elem.text.strip()
                if not text:
                    continue

                text = sanitize_document_text(text)
                style_name = p_elem.style.name.lower() if p_elem.style else ""

                if "heading" in style_name or "title" in style_name:
                    current_section = text
                    continue

                total_chars += len(text)
                estimated_page = max(1, total_chars // 2500 + 1)
                full_text_parts.append(text)

                char_start = global_char_offset
                char_end = char_start + len(text)
                chunks.append(
                    ParsedChunk(
                        content=text,
                        page_number=estimated_page,
                        section_title=current_section,
                        chunk_index=chunk_idx,
                        token_count=len(text.split()),
                        char_offset_start=char_start,
                        char_offset_end=char_end,
                        source_type="PARAGRAPH",
                        metadata={"style": style_name, "source_file": filename},
                    )
                )
                chunk_idx += 1
                global_char_offset = char_end + 1

            # Table
            elif element.tag.endswith("tbl"):
                t_elem = None
                for t in doc.tables:
                    if t._element == element:
                        t_elem = t
                        break
                if not t_elem:
                    continue

                table_count += 1
                table_rows: List[str] = []
                for row in t_elem.rows:
                    row_cells = [sanitize_document_text(cell.text.strip()) for cell in row.cells]
                    table_rows.append(" | ".join(row_cells))

                if table_rows:
                    table_md = "\n".join(table_rows)
                    total_chars += len(table_md)
                    estimated_page = max(1, total_chars // 2500 + 1)
                    full_text_parts.append(table_md)

                    char_start = global_char_offset
                    char_end = char_start + len(table_md)
                    chunks.append(
                        ParsedChunk(
                            content=table_md,
                            page_number=estimated_page,
                            section_title=f"{current_section} (Table {table_count})",
                            chunk_index=chunk_idx,
                            token_count=len(table_md.split()),
                            char_offset_start=char_start,
                            char_offset_end=char_end,
                            source_type="TABLE",
                            metadata={"table_index": table_count, "source_file": filename},
                        )
                    )
                    chunk_idx += 1
                    global_char_offset = char_end + 1

        raw_text = "\n\n".join(full_text_parts)
        estimated_pages = max(1, total_chars // 2500 + 1)

        return ParsedDocument(
            raw_text=raw_text,
            page_count=estimated_pages,
            table_count=table_count,
            chunks=chunks,
            metadata={"filename": filename, "format": "DOCX", "table_count": table_count},
        )
